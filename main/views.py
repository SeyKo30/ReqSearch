from django.http import HttpResponse, HttpResponseNotFound, HttpResponseForbidden
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import logout
from django.contrib.auth.decorators import login_required
from .forms import TextInputForm, SignUpForm, CompanyForm, DocumentForm
from .models import Company
from docx import Document
import re


@login_required
def profile(request):
    user = request.user
    companies = Company.objects.filter(user=user)
    return render(request, 'main/profile.html', {'user': user, 'companies': companies})


@login_required
def create_company(request):
    if request.method == 'POST':
        form = CompanyForm(request.POST)
        if form.is_valid():
            company = form.save(commit=False)
            company.user = request.user  
            company.save()
            return redirect('success_page')
    else:
        form = CompanyForm()

    return render(request, 'main/company_form.html', {'form': form})


def success_page(request):
    return render(request, 'main/success.html')


@login_required
def company_list(request):
    companies = Company.objects.all()
    return render(request, 'main/company_list.html', {'companies': companies})


def logout_view(request):
    logout(request)
    return redirect('index')


@login_required
def profile_view(request):
    user = request.user
    companies = Company.objects.filter(user=user) 

    context = {
        'user': user,
        'companies': companies,  
    }
    return render(request, 'main/profile.html', context)


def signup(request):
    if request.method == 'POST':
        form = SignUpForm(request.POST)
        if form.is_valid():
            form.save()
            return redirect('login')  
    else:
        form = SignUpForm()
    return render(request, 'main/signup.html', {'form': form})


def about_us(request):
    return render(request, 'main/about_us.html')


def processing_page(request):
    sentences_with_keywords1 = request.session.get('sentences_with_keywords1', [])
    sentences_with_keywords2 = request.session.get('sentences_with_keywords2', [])

    context = {
        'sentences_with_keywords1': sentences_with_keywords1,
        'sentences_with_keywords2': sentences_with_keywords2,
    }

    return render(request, 'main/processing_page.html', context)


def delete_company(request, company_id):
    company = get_object_or_404(Company, id=company_id)
    
    if company.user != request.user:
        return HttpResponseForbidden("You don't have permission to delete this company.")

    company.delete()
    return redirect('profile')


def upload_document(request):
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)

        if form.is_valid():
            keywords_list1 = ['довідку', 'гарантійний', 'гарантійного', 'Згода', 'Згоду', 'Лист',
                              'довідка', 'довільной', 'довільній'
                              ]

            exclude_words_list1 = ['не містить вихідного номера', 'Вимога щодо засвідчення', 'Наприклад',
                                   'Відгук', 'Повноваження щодо підпису', 'банківської', 'не повинен складати',
                                   'який надав найбільш економічно', 'Приклади формальних помилок',
                                   'не передбачений (необов’язковий) для учасника', 'аномально низька ціна визначається',
                                   'гарантійного випадку', 'прайс-лист', 'гарантійний талон',
                                   'Наприклад', 'Наприклад', 'Наприклад', 'Наприклад', 'Наприклад',
                                   'У випадку усунення Виконавцем недоліків', 'Учасником надано', 'Учасником не надано ',
                                   'гарантійного терміну', 'гарантійний термін', 'гарантійний строк', 'гарантійного строку', 'test', 'test', 'test', 'test'
                                   ]

            keywords_list2 = ['Витяг', 'Витягу', 'Паспорт', 'Сертифікат', 'ліцензія', 'Наказ', 'протокол', 'Договір', 'Договори',
                              'акти', 'Накладна', 'Свідоцтва', 'Свідоцтво', 'Виданий', 'Датований', 'Відповідним', 'Відповідною',
                              'Накладну', 'Накладні', 'книжки', 'книжку', 'книжка'
                              ]
            exclude_words_list2 = ['Про затвердження Переліку формальних помилок', 'повідомлення про намір',
                                   'Протокол розкриття тендерних пропозицій', 'У разі відхилення тендерної',
                                   'укладається відповідно до', 'Замовник укладає договір про закупівлю з учасником',
                                   'активи в установленому', 'активи такої особи', 'Про електронні довірчі послуги', 'test',
                                   'Учасником надано', 'Учасником не надано ' 'гарантійного строку', 'test', 'test', 'test', 'test', 'Гарантійний строк',
                                   'наказом', 'Протокол Уповноваженої особи', 'інформацією про маркування', 'вважається договір'
                                   ]

            uploaded_file = form.save()
            doc = Document(uploaded_file.file)

            sentences = process_document_text(doc)

            sentences_with_keywords1 = filter_sentences(sentences, exclude_words_list1, keywords_list1)
            sentences_with_keywords2 = filter_sentences(sentences, exclude_words_list2, keywords_list2)

            request.session['sentences_with_keywords1'] = sentences_with_keywords1
            request.session['sentences_with_keywords2'] = sentences_with_keywords2

            return redirect('processing_page')

    else:
        form = DocumentForm()

    return render(request, 'main/upload.html', {'form': form})


def process_document_text(doc):
    # Извлечение текста из абзацев
    paragraphs_text = ' '.join(para.text if para.text is not None else '' for para in doc.paragraphs)

    # Извлечение текста из других частей документа (например, заголовков)
    other_parts_text = ''
    for part in doc.element.body:
        if part.tag.endswith('tbl'):  # Обработка таблиц
            for row in part:
                for cell in row:
                    other_parts_text += ' '.join(cell.text if cell.text is not None else '' for cell in cell)
        elif part.tag.endswith('hdr'):  # Обработка заголовков
            other_parts_text += ' '.join(para.text if para.text is not None else '' for para in part.paragraphs)
    
    # Объединение текста из абзацев и других частей документа
    full_text = paragraphs_text + ' ' + other_parts_text
    
    # Разбиение текста на предложения
    sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?|;|\n|\r|\x85)\s', full_text)

    # Возвращаем список предложений
    return sentences


def filter_sentences(sentences, exclude_words, keywords):
    sentences_with_keywords = []

    for sentence in sentences:
        has_exclude_word = any(exclude_word.lower() in sentence.lower() for exclude_word in exclude_words)

        if not has_exclude_word:
            for keyword in keywords:
                if keyword.lower() in sentence.lower():
                    sentences_with_keywords.append(sentence)
                    break

    return sentences_with_keywords


def index(request):
    return render(request, 'main/index.html')


def page_not_found(request, exception):
    return HttpResponseNotFound('<h1>Такой страницы не существует. <a href="http://127.0.0.1:8000" rel="noopener" target="_blank" >вернуться на главную</a></h1>')
